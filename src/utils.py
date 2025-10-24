# src/utils.py
"""
Enhanced Utility Module for the Plagiarism Checker Project.
Robust PDF processing with multiple fallback methods and comprehensive error handling.
"""

import os
import docx
from PyPDF2 import PdfReader
from io import BytesIO
import sqlite3
import streamlit as st
import re
import logging
import fitz  # PyMuPDF
import pdfplumber
from typing import Optional, Tuple, Dict, Any
import tempfile

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


def read_txt_file(file_input) -> Optional[str]:
    """
    Read text file with robust encoding detection and error handling.
    
    Args:
        file_input: File path or BytesIO object
        
    Returns:
        Extracted text or None if failed
    """
    encodings = ['utf-8', 'latin-1', 'windows-1252', 'iso-8859-1', 'cp1252']
    
    try:
        if isinstance(file_input, str):
            # File path
            for encoding in encodings:
                try:
                    with open(file_input, 'r', encoding=encoding) as file:
                        content = file.read()
                        if content.strip():
                            logger.info(f"Successfully read TXT with {encoding}")
                            return clean_academic_text(content)
                except UnicodeDecodeError:
                    continue
        else:
            # BytesIO object
            file_input.seek(0)
            raw_data = file_input.read()
            for encoding in encodings:
                try:
                    content = raw_data.decode(encoding)
                    if content.strip():
                        logger.info(f"Successfully decoded TXT with {encoding}")
                        return clean_academic_text(content)
                except UnicodeDecodeError:
                    continue
        
        logger.error("Failed to decode TXT file with any encoding")
        return None
        
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        return None

def read_docx_file(file_input) -> Optional[str]:
    """
    Read DOCX file from path or BytesIO with comprehensive error handling.
    
    Args:
        file_input: File path or BytesIO object
        
    Returns:
        Extracted text or None if failed
    """
    try:
        if isinstance(file_input, str):
            document = docx.Document(file_input)
        else:
            file_input.seek(0)
            document = docx.Document(file_input)
        
        full_text = []
        for para in document.paragraphs:
            text = para.text.strip()
            # Skip empty paragraphs and very short ones (likely formatting)
            if text and len(text) > 3 and not text.isspace():
                full_text.append(text)
        
        if not full_text:
            logger.warning("DOCX file contains no extractable text")
            return None
        
        result = '\n'.join(full_text)
        logger.info(f"Successfully extracted {len(full_text)} paragraphs from DOCX")
        return clean_academic_text(result)
        
    except Exception as e:
        logger.error(f"Error reading DOCX file: {e}")
        return None

def read_pdf_file(file_input) -> Optional[str]:
    """
    Simple PDF reading using PyMuPDF as primary method.
    """
    try:
        if hasattr(file_input, 'read'):
            file_input.seek(0)
            doc = fitz.open(stream=file_input.read(), filetype="pdf")
        else:
            doc = fitz.open(file_input)
        
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text and text.strip():
                text_parts.append(text.strip())
        
        doc.close()
        
        if text_parts:
            full_text = " ".join(text_parts)
            # Basic cleaning
            full_text = re.sub(r'\s+', ' ', full_text)
            full_text = full_text.strip()
            
            if len(full_text.split()) >= 50:  # Minimum word count
                logger.info(f"PDF extraction successful: {len(full_text.split())} words")
                return full_text
        
        logger.error("PDF extraction failed - insufficient text extracted")
        return None
        
    except Exception as e:
        logger.error(f"PDF reading failed: {e}")
        return None

def clean_academic_text(text: str) -> str:
    """
    Comprehensive cleaning of academic text with enhanced patterns.
    
    Args:
        text: Raw text to clean
        
    Returns:
        Cleaned academic text
    """
    if not text:
        return ""
    
    # Enhanced academic paper cleaning patterns
    patterns_to_remove = [
        r'arXiv:\d+\.\d+v\d+\[.*?\]\s+\d+.*?\d{4}',  # arXiv headers
        r'Proceedings of.*?(\n|$)',
        r'Journal of.*?(\n|$)',
        r'Conference on.*?(\n|$)',
        r'Copyright ©.*?(\n|$)',
        r'©.*?(\n|$)',
        r'\bdoi:.*?\b',
        r'https?://\S+',
        r'\S+@\S+\.\S+',  # Emails
        r'\n\s*\d+\s*\n',  # Standalone page numbers
        r'^\\d+$',  # Lines with only numbers
        r'CONFIDENTIAL|DRAFT|PREPRINT|SUBMITTED',  # Common document labels
        r'\[.*?\d+.*?\]',  # Citation markers like [1], [2-5]
    ]
    
    for pattern in patterns_to_remove:
        text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)
    
    # Remove excessive whitespace and normalize
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'\n+', '\n', text)
    
    return text.strip()

def read_uploaded_file(uploaded_file) -> Optional[str]:
    """
    Read uploaded file with comprehensive error handling and format detection.
    
    Args:
        uploaded_file: Streamlit uploaded file object
        
    Returns:
        Extracted text or None if failed
    """
    try:
        if not uploaded_file or uploaded_file.size == 0:
            logger.error("Uploaded file is empty or invalid")
            return None
        
        filename = uploaded_file.name.lower()
        file_size_mb = uploaded_file.size / (1024 * 1024)
        
        logger.info(f"Processing uploaded file: {filename} ({file_size_mb:.2f} MB)")
        
        # Check file size limits
        if file_size_mb > 50:  # 50MB limit
            logger.error(f"File too large: {file_size_mb:.2f} MB")
            return None
        
        # Dispatch to appropriate reader based on file extension
        if filename.endswith('.txt'):
            text = read_txt_file(BytesIO(uploaded_file.getvalue()))
            return text  # Already cleaned in read_txt_file
            
        elif filename.endswith('.docx'):
            text = read_docx_file(BytesIO(uploaded_file.getvalue()))
            return text  # Already cleaned in read_docx_file
            
        elif filename.endswith('.pdf'):
            text = read_pdf_file(BytesIO(uploaded_file.getvalue()))
            return text  # Already cleaned in PDF processor
            
        else:
            logger.error(f"Unsupported file type: {filename}")
            return None
            
    except Exception as e:
        logger.error(f"Error processing uploaded file {uploaded_file.name}: {e}")
        return None

def validate_document_text(text: str, filename: str, min_words: int = 100) -> Tuple[bool, str]:
    """
    Comprehensive document validation with detailed error reporting.
    
    Args:
        text: Extracted text to validate
        filename: Original filename for error messages
        min_words: Minimum number of words required
        
    Returns:
        Tuple of (is_valid, error_message)
    """
    if not text:
        error_msg = f"""
❌ **No text extracted from '{filename}'**
        
**Possible issues:**
• The file may be encrypted or password-protected
• It could be image-based (scanned document)
• The file might be corrupted or in an unsupported format
• The document may contain no extractable text content

**Suggested solutions:**
• Try a different PDF file format
• Use text-based PDFs instead of scanned documents
• Ensure the document is not password-protected
• Convert image-based PDFs to text using OCR software
"""
        return False, error_msg
    
    words = text.split()
    word_count = len(words)
    
    if word_count < min_words:
        error_msg = f"""
❌ **Insufficient text in '{filename}'**

**Details:**
• Extracted only **{word_count} words** (minimum {min_words} required)
• This often happens with image-based PDFs or scanned documents
• The document may contain mostly images, tables, or complex formatting

**Suggested solutions:**
• Use text-based PDFs instead of scanned documents
• Try converting the document to a different format (DOCX, TXT)
• For scanned PDFs, use OCR software to extract text first
• Check if the document has selectable text (try copying text manually)
"""
        return False, error_msg
    
    # Check text quality metrics
    valid_words = [word for word in words if len(word) > 1]
    valid_ratio = len(valid_words) / len(words) if words else 0
    
    if valid_ratio < 0.7:  # At least 70% valid words
        error_msg = f"""
❌ **Poor text quality in '{filename}'**

**Details:**
• Only {valid_ratio:.1%} of extracted words appear valid
• This suggests garbled text, encoding issues, or extraction problems
• Common with poorly scanned documents or complex layouts

**Suggested solutions:**
• Use a higher quality source document
• Try OCR software for scanned documents
• Convert to a different file format
"""
        return False, error_msg
    
    # Check for reasonable sentence structure
    sentences = re.split(r'[.!?]+', text)
    valid_sentences = [s for s in sentences if len(s.split()) > 3]
    
    if len(valid_sentences) < 3:
        error_msg = f"""
❌ **Poor document structure in '{filename}'**

**Details:**
• Insufficient sentence structure detected
• Only {len(valid_sentences)} valid sentences found (minimum 3 required)
• This may indicate extraction issues or non-text content

**Suggested solutions:**
• Verify the document contains readable text content
• Try a different file format or source
• Check if text can be manually selected and copied
"""
        return False, error_msg
    
    return True, ""

def extract_metadata_from_text(text: str, filename: str) -> Dict[str, Any]:
    """
    Extract comprehensive metadata from academic paper text.
    
    Args:
        text: Document text content
        filename: Original filename
        
    Returns:
        Dictionary containing metadata
    """
    if not text:
        return {
            'filename': filename,
            'word_count': 0,
            'estimated_pages': 0,
            'has_arxiv_ref': False,
            'has_doi': False,
            'has_citations': False,
            'possible_title': filename,
            'status': 'No text extracted',
            'quality_score': 0
        }
    
    words = text.split()
    sentences = re.split(r'[.!?]+', text)
    
    metadata = {
        'filename': filename,
        'word_count': len(words),
        'estimated_pages': max(1, len(text) // 2500),
        'sentence_count': len([s for s in sentences if len(s.split()) > 3]),
        'has_arxiv_ref': bool(re.search(r'arxiv:\d+\.\d+', text.lower())),
        'has_doi': bool(re.search(r'doi:\s*\S+', text.lower()) or 'doi.org' in text.lower()),
        'has_citations': bool(re.search(r'\[\s*\d+\s*\]', text) or re.search(r'\(\s*\w+,\s*\d{4}\s*\)', text)),
        'status': 'Success',
        'quality_score': min(100, len(words) // 10)  # Simple quality metric
    }
    
    # Try to extract title (first reasonable length line)
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        word_count = len(line.split())
        if 5 <= word_count <= 20 and len(line) < 200:  # Reasonable title characteristics
            metadata['possible_title'] = line
            break
    
    if 'possible_title' not in metadata:
        # Use first 200 chars as fallback
        metadata['possible_title'] = text[:200].split('\n')[0] + '...' if len(text) > 200 else text
    
    # Detect academic paper characteristics
    academic_indicators = [
        'abstract', 'introduction', 'methodology', 'results', 'conclusion',
        'references', 'bibliography', 'acknowledgements'
    ]
    
    metadata['academic_indicators'] = [
        indicator for indicator in academic_indicators 
        if indicator in text.lower()
    ]
    
    metadata['is_likely_academic'] = len(metadata['academic_indicators']) >= 2
    
    return metadata

@st.cache_data
def get_all_documents_from_db(db_path: str = "corpus.db") -> list:
    """
    Retrieve all documents from the SQLite database with error handling.
    
    Args:
        db_path: Path to the SQLite database file
        
    Returns:
        List of document tuples (filename, document_vector)
    """
    if not os.path.exists(db_path):
        logger.warning(f"Database file not found at {db_path}")
        return []

    conn = None
    try:
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        # Check if table exists
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table' AND name='documents'")
        if not cursor.fetchone():
            logger.warning("Documents table does not exist in database")
            return []

        select_query = "SELECT filename, document_vector FROM documents WHERE document_vector IS NOT NULL"
        cursor.execute(select_query)

        documents = cursor.fetchall()
        logger.info(f"Retrieved {len(documents)} documents from database")
        return documents
    
    except sqlite3.Error as e:
        logger.error(f"Database error: {e}")
        return []
    finally:
        if conn:
            conn.close()

def get_file_type_info(filename: str) -> Dict[str, str]:
    """
    Get information about supported file types.
    
    Args:
        filename: The filename to analyze
        
    Returns:
        Dictionary with file type information
    """
    file_ext = os.path.splitext(filename.lower())[1]
    
    info = {
        '.pdf': {
            'name': 'PDF Document',
            'support': 'Excellent (text-based), Limited (scanned)',
            'limitations': 'Scanned PDFs require OCR',
            'recommendation': 'Use text-selectable PDFs for best results'
        },
        '.docx': {
            'name': 'Word Document', 
            'support': 'Excellent',
            'limitations': 'Complex formatting may affect extraction',
            'recommendation': 'Ideal format for text extraction'
        },
        '.txt': {
            'name': 'Text File',
            'support': 'Perfect', 
            'limitations': 'None',
            'recommendation': 'Best format for accurate analysis'
        }
    }
    
    return info.get(file_ext, {
        'name': 'Unknown Format',
        'support': 'Not supported',
        'limitations': 'Unsupported file type',
        'recommendation': 'Convert to PDF, DOCX, or TXT'
    })
